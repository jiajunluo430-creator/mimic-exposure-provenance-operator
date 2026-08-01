from __future__ import annotations

import hashlib
import io
import json
import re
import sys
import tarfile
import time
import urllib.error
import urllib.parse
import urllib.request
import xml.etree.ElementTree as ET
import zipfile
from datetime import datetime
from pathlib import Path

import pandas as pd
from docx import Document
from openpyxl import load_workbook
from pypdf import PdfReader


PROJECT = Path(__file__).resolve().parents[1]
PRIOR = PROJECT / "outputs" / "jamia_prereview_upgrade_v1_0"
SAMPLE = PRIOR / "tables" / "published_operator_landscape_sample.csv"
MAIN_XML = PRIOR / "literature_raw"
OUTPUT = PROJECT / "outputs" / "jamia_residual_provenance_v1_0"
TABLES = OUTPUT / "tables"
LOGS = OUTPUT / "logs"
MANIFESTS = OUTPUT / "manifests"
RAW = OUTPUT / "literature_scope_raw"
SUPP_RAW = RAW / "supplements"
REPO_RAW = RAW / "repositories"
OA_RAW = RAW / "oa_packages"
EUROPE_PMC_RAW = RAW / "europe_pmc_supplement_bundles"
TEXT = OUTPUT / "literature_scope_text"
ENVIRONMENT = PROJECT / "environment"
ADDENDUM = (
    PROJECT
    / "contracts"
    / "jamia_residual_provenance_addendum_v1.0_2026-07-31.md"
)
CODEBOOK = (
    PROJECT
    / "contracts"
    / "published_operator_landscape_codebook_v1.0_2026-07-31.md"
)
CODEBOOK_FIELDS = (
    PROJECT
    / "contracts"
    / "published_operator_landscape_codebook_fields_v1.0.csv"
)

EXPECTED_ADDENDUM_SHA256 = (
    "af533e4d3a9b636c368dc6c76cc3e3ea472c77f9884476520d341ae09575dcfd"
)
EXPECTED_CODEBOOK_SHA256 = (
    "ecfab4f17c4c277c3cf0a26e4aecb4a9e6fae1193149ac687c599baf717bea53"
)
EXPECTED_CODEBOOK_FIELDS_SHA256 = (
    "d2ed9c4e01ba1477d32690a4a1e5e62026b6d497e5f35c414154d1ea2d64279d"
)
EXPECTED_SAMPLE_SHA256 = (
    "69fc62e3ff0bf3f8f9eb626817a6342ebcfebf46f9ca11d7a7ed07418e407678"
)
USER_AGENT = "N1-MIMIC-medication-provenance-scope-audit/1.0"
XLINK = "{http://www.w3.org/1999/xlink}href"

for directory in (
    TABLES,
    LOGS,
    MANIFESTS,
    SUPP_RAW,
    REPO_RAW,
    OA_RAW,
    EUROPE_PMC_RAW,
    TEXT,
    ENVIRONMENT,
):
    directory.mkdir(parents=True, exist_ok=True)

LOG_PATH = LOGS / "37_audit_published_operator_evidence_scope.log"


def log(message: str) -> None:
    line = f"{datetime.now().astimezone().isoformat()}\t{message}"
    print(line, flush=True)
    with LOG_PATH.open("a", encoding="utf-8") as handle:
        handle.write(line + "\n")


def sha256(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as handle:
        for block in iter(lambda: handle.read(1024 * 1024), b""):
            digest.update(block)
    return digest.hexdigest()


def request_bytes(url: str, attempts: int = 4) -> tuple[int, bytes, str]:
    for attempt in range(1, attempts + 1):
        request = urllib.request.Request(url, headers={"User-Agent": USER_AGENT})
        try:
            with urllib.request.urlopen(request, timeout=90) as response:
                return int(response.status), response.read(), response.geturl()
        except urllib.error.HTTPError as error:
            if error.code in (429, 500, 502, 503, 504) and attempt < attempts:
                time.sleep(min(30, 2**attempt))
                continue
            return int(error.code), str(error).encode(), url
        except Exception as error:
            if attempt < attempts:
                time.sleep(min(20, 2**attempt))
                continue
            return 0, repr(error).encode(), url
    raise AssertionError("unreachable")


def oa_package_supplements(
    pmcid: str, hrefs: list[str]
) -> tuple[dict[str, bytes], str, int]:
    if not hrefs:
        return {}, "", 0
    oa_url = f"https://www.ncbi.nlm.nih.gov/pmc/utils/oa/oa.fcgi?id={pmcid}"
    status, metadata, _ = request_bytes(oa_url)
    if status != 200:
        return {}, oa_url, status
    try:
        root = ET.fromstring(metadata)
        link = next(
            node.attrib.get("href", "")
            for node in root.findall(".//link")
            if node.attrib.get("format") == "tgz"
        )
    except (ET.ParseError, StopIteration):
        return {}, oa_url, 404
    package_url = re.sub(r"^ftp://", "https://", link)
    package_path = OA_RAW / f"{pmcid}.tar.gz"
    if package_path.exists() and package_path.stat().st_size > 0:
        package = package_path.read_bytes()
        package_status = 200
    else:
        package_status, package, package_url = request_bytes(package_url)
        if package_status == 200:
            package_path.write_bytes(package)
    if package_status != 200:
        return {}, package_url, package_status
    wanted = {safe_name(href).lower(): safe_name(href) for href in hrefs}
    found: dict[str, bytes] = {}
    try:
        with tarfile.open(fileobj=io.BytesIO(package), mode="r:gz") as archive:
            for member in archive.getmembers():
                if not member.isfile():
                    continue
                basename = Path(member.name).name
                expected = wanted.get(basename.lower())
                if expected is None:
                    continue
                handle = archive.extractfile(member)
                if handle is not None:
                    found[expected] = handle.read()
    except tarfile.TarError:
        return {}, package_url, 422
    return found, package_url, package_status


def europe_pmc_supplements(
    pmcid: str, hrefs: list[str]
) -> tuple[dict[str, bytes], str, int]:
    if not hrefs:
        return {}, "", 0
    url = (
        "https://www.ebi.ac.uk/europepmc/webservices/rest/"
        f"{pmcid}/supplementaryFiles"
    )
    bundle_path = EUROPE_PMC_RAW / f"{pmcid}.zip"
    if bundle_path.exists() and bundle_path.stat().st_size > 0:
        payload = bundle_path.read_bytes()
        status = 200
    else:
        status, payload, url = request_bytes(url)
        if status == 200:
            bundle_path.write_bytes(payload)
    if status != 200:
        return {}, url, status
    wanted = {safe_name(href).lower(): safe_name(href) for href in hrefs}
    found: dict[str, bytes] = {}
    try:
        with zipfile.ZipFile(io.BytesIO(payload)) as archive:
            for member in archive.infolist():
                if member.is_dir():
                    continue
                basename = Path(member.filename).name
                expected = wanted.get(basename.lower())
                if expected is not None:
                    found[expected] = archive.read(member)
    except zipfile.BadZipFile:
        return {}, url, 422
    return found, url, status


def safe_name(value: str) -> str:
    name = Path(urllib.parse.urlparse(value).path).name or "supplement.bin"
    return re.sub(r"[^A-Za-z0-9._-]+", "_", name)


def xml_text(data: bytes) -> str:
    root = ET.fromstring(data)
    return " ".join("".join(root.itertext()).split())


def extract_docx(data: bytes) -> str:
    document = Document(io.BytesIO(data))
    chunks = [paragraph.text for paragraph in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            chunks.append("\t".join(cell.text for cell in row.cells))
    return "\n".join(chunks)


def extract_xlsx(data: bytes) -> str:
    workbook = load_workbook(io.BytesIO(data), read_only=True, data_only=True)
    chunks: list[str] = []
    for sheet in workbook.worksheets:
        chunks.append(f"[SHEET {sheet.title}]")
        for row in sheet.iter_rows(values_only=True):
            values = ["" if value is None else str(value) for value in row]
            if any(values):
                chunks.append("\t".join(values))
    return "\n".join(chunks)


def extract_pdf(data: bytes) -> str:
    reader = PdfReader(io.BytesIO(data))
    return "\n".join(page.extract_text() or "" for page in reader.pages)


def extract_zip(data: bytes, depth: int = 0) -> str:
    if depth > 2:
        return ""
    chunks: list[str] = []
    with zipfile.ZipFile(io.BytesIO(data)) as archive:
        for member in archive.infolist():
            if member.is_dir() or member.file_size > 50_000_000:
                continue
            child = archive.read(member)
            text, status = extract_content(member.filename, child, depth + 1)
            if text:
                chunks.append(f"\n[ARCHIVE MEMBER {member.filename}; {status}]\n{text}")
    return "\n".join(chunks)


def extract_content(name: str, data: bytes, depth: int = 0) -> tuple[str, str]:
    suffix = Path(name).suffix.lower()
    try:
        if suffix == ".docx":
            return extract_docx(data), "docx_extracted"
        if suffix in {".xlsx", ".xlsm"}:
            return extract_xlsx(data), "xlsx_extracted"
        if suffix == ".pdf":
            return extract_pdf(data), "pdf_extracted"
        if suffix == ".zip":
            return extract_zip(data, depth), "zip_extracted"
        if suffix in {
            ".txt",
            ".md",
            ".csv",
            ".tsv",
            ".sql",
            ".py",
            ".r",
            ".json",
            ".yaml",
            ".yml",
            ".xml",
        }:
            return data.decode("utf-8", errors="replace"), "text_extracted"
        return "", "binary_not_text_reviewable"
    except Exception as error:
        return "", f"extraction_error:{type(error).__name__}:{error}"


def repo_links(text: str) -> list[str]:
    links = re.findall(
        r"https?://(?:www\.)?(?:github\.com|gitlab\.com|osf\.io|"
        r"zenodo\.org|figshare\.com)/[^\s<>\"']+",
        text.replace("&amp;", "&"),
        flags=re.IGNORECASE,
    )
    cleaned: set[str] = set()
    for link in links:
        link = link.rstrip(".,;:)]}")
        link = re.sub(r"github\.com/+", "github.com/", link, flags=re.I)
        cleaned.add(link)
    return sorted(cleaned)


def github_repo(link: str) -> tuple[str, str] | None:
    parsed = urllib.parse.urlparse(link)
    if parsed.netloc.lower() not in {"github.com", "www.github.com"}:
        return None
    parts = [part for part in parsed.path.split("/") if part]
    if len(parts) < 2:
        return None
    return parts[0], parts[1].removesuffix(".git")


def is_generic_repo(link: str) -> bool:
    parsed_url = urllib.parse.urlparse(link)
    path_parts = [part for part in parsed_url.path.split("/") if part]
    if (
        parsed_url.netloc.lower() in {"github.com", "www.github.com"}
        and len(path_parts) == 1
        and path_parts[0].lower() == "mit-lcp"
    ):
        return True
    parsed = github_repo(link)
    if parsed is None:
        return False
    owner, repo = (value.lower() for value in parsed)
    return owner == "mit-lcp" and (
        repo.startswith("mimic") or repo in {"", "mimic-code"}
    )


def keyword_hits(text: str) -> dict[str, int]:
    patterns = {
        "native_table": r"\b(?:emar|prescriptions|pharmacy|poe|inputevents|chartevents)\b",
        "identity_field": r"\b(?:poe_id|pharmacy_id|emar_id|itemid|ndc|gsn|formulary_drug_cd)\b",
        "event_semantics": r"\b(?:event_txt|not given|held|hold dose|refused|flushed|confirmed|given)\b",
        "time_rule": r"\b(?:within|first)\s+\d+\s*(?:h|hour|hours|day|days)\b|\bicu admission\b",
        "dose_route": r"\b(?:dose|route|subcutaneous|intravenous|oral|po|iv|sc|sq)\b",
        "executable_code": r"\b(?:select|from|where|join|regexp|str_detect|read_csv|duckdb)\b",
    }
    return {
        label: len(re.findall(pattern, text, flags=re.IGNORECASE))
        for label, pattern in patterns.items()
    }


def main() -> None:
    hashes = {
        "addendum": (ADDENDUM, EXPECTED_ADDENDUM_SHA256),
        "codebook": (CODEBOOK, EXPECTED_CODEBOOK_SHA256),
        "codebook_fields": (CODEBOOK_FIELDS, EXPECTED_CODEBOOK_FIELDS_SHA256),
        "sample": (SAMPLE, EXPECTED_SAMPLE_SHA256),
    }
    for label, (path, expected) in hashes.items():
        observed = sha256(path)
        if observed != expected:
            raise RuntimeError(f"{label} hash mismatch: {observed}")
    log("PASS frozen addendum, codebook, and 40-study sample hashes")

    sample = pd.read_csv(SAMPLE, dtype=str)
    if sample.shape[0] != 40 or sample["pmcid"].nunique() != 40:
        raise RuntimeError("Frozen sample cardinality failed")

    supplement_rows: list[dict[str, object]] = []
    repo_by_pmc: dict[str, set[str]] = {}
    text_by_pmc: dict[str, list[str]] = {}
    linked_supplements_by_pmc: dict[str, int] = {}
    retrieved_supplements_by_pmc: dict[str, int] = {}
    reviewable_supplements_by_pmc: dict[str, int] = {}

    for _, article in sample.iterrows():
        pmcid = str(article["pmcid"])
        xml_path = MAIN_XML / f"{pmcid}.xml"
        data = xml_path.read_bytes()
        root = ET.fromstring(data)
        repo_by_pmc[pmcid] = set(repo_links(data.decode("utf-8", errors="replace")))
        text_by_pmc[pmcid] = []
        hrefs: list[str] = []
        for node in root.findall(".//supplementary-material"):
            for descendant in node.iter():
                href = descendant.attrib.get(XLINK, "").strip()
                if href:
                    hrefs.append(href)
        hrefs = list(dict.fromkeys(hrefs))
        linked_supplements_by_pmc[pmcid] = len(hrefs)
        retrieved_supplements_by_pmc[pmcid] = 0
        reviewable_supplements_by_pmc[pmcid] = 0
        article_dir = SUPP_RAW / pmcid
        article_dir.mkdir(parents=True, exist_ok=True)
        europe_payloads, europe_url, europe_status = europe_pmc_supplements(
            pmcid, hrefs
        )
        oa_payloads: dict[str, bytes] = {}
        oa_package_url = ""
        oa_package_status = 0
        if len(europe_payloads) < len(hrefs):
            oa_payloads, oa_package_url, oa_package_status = oa_package_supplements(
                pmcid, hrefs
            )

        for href in hrefs:
            filename = safe_name(href)
            target = article_dir / filename
            direct_url = (
                f"https://pmc.ncbi.nlm.nih.gov/articles/{pmcid}/bin/"
                + urllib.parse.quote(filename)
            )
            if filename in europe_payloads:
                status, payload, final_url = 200, europe_payloads[filename], europe_url
                target.write_bytes(payload)
                retrieval_source = "europe_pmc_supplementary_files"
            elif filename in oa_payloads:
                status, payload, final_url = 200, oa_payloads[filename], oa_package_url
                target.write_bytes(payload)
                retrieval_source = "ncbi_oa_package"
            else:
                status, payload, final_url = request_bytes(direct_url)
                retrieval_source = "pmc_direct_fallback"
                if status == 200 and not payload.lstrip().lower().startswith(b"<!doctype html"):
                    target.write_bytes(payload)
                elif status == 200:
                    status = 403
                time.sleep(0.05)
            extracted, extraction_status = (
                extract_content(filename, payload) if status == 200 else ("", "not_retrieved")
            )
            if status == 200:
                retrieved_supplements_by_pmc[pmcid] += 1
            if extracted.strip():
                reviewable_supplements_by_pmc[pmcid] += 1
                text_path = TEXT / f"{pmcid}__{filename}.txt"
                text_path.write_text(extracted, encoding="utf-8")
                text_by_pmc[pmcid].append(extracted)
                repo_by_pmc[pmcid].update(repo_links(extracted))
            hits = keyword_hits(extracted)
            supplement_rows.append(
                {
                    "pmcid": pmcid,
                    "supplement_href": href,
                    "request_url": direct_url,
                    "final_url": final_url,
                    "retrieval_source": retrieval_source,
                    "europe_pmc_status": europe_status,
                    "oa_package_status": oa_package_status,
                    "http_status": status,
                    "bytes": len(payload) if status == 200 else 0,
                    "sha256": hashlib.sha256(payload).hexdigest() if status == 200 else "",
                    "extraction_status": extraction_status,
                    "reviewable_text_characters_n": len(extracted),
                    **hits,
                }
            )
        log(
            f"CHECKPOINT {pmcid} supplements={len(hrefs)} "
            f"retrieved={retrieved_supplements_by_pmc[pmcid]} "
            f"reviewable={reviewable_supplements_by_pmc[pmcid]}"
        )

    repo_rows: list[dict[str, object]] = []
    for pmcid, links in sorted(repo_by_pmc.items()):
        for link in sorted(links):
            generic = is_generic_repo(link)
            parsed = github_repo(link)
            status = 0
            final_url = link
            payload = b""
            extracted = ""
            extraction_status = "not_retrieved"
            if parsed and not generic:
                owner, repo = parsed
                api_url = f"https://api.github.com/repos/{owner}/{repo}/zipball"
                target = REPO_RAW / f"{pmcid}__{owner}__{repo}.zip"
                if target.exists() and target.stat().st_size > 0:
                    status, payload, final_url = 200, target.read_bytes(), api_url
                else:
                    status, payload, final_url = request_bytes(api_url)
                    if status == 200:
                        target.write_bytes(payload)
                    time.sleep(0.2)
                if status == 200:
                    extracted, extraction_status = extract_content(target.name, payload)
                    if extracted.strip():
                        (TEXT / f"{pmcid}__repo__{owner}__{repo}.txt").write_text(
                            extracted, encoding="utf-8"
                        )
            elif generic:
                extraction_status = "generic_repository_not_article_specific"
            else:
                status, payload, final_url = request_bytes(link)
                if status == 200:
                    extracted = payload.decode("utf-8", errors="replace")
                    extraction_status = "landing_page_extracted"
            repo_rows.append(
                {
                    "pmcid": pmcid,
                    "repository_url": link,
                    "generic_repository": generic,
                    "http_status": status,
                    "final_url": final_url,
                    "bytes": len(payload) if status == 200 else 0,
                    "sha256": hashlib.sha256(payload).hexdigest() if status == 200 else "",
                    "extraction_status": extraction_status,
                    "reviewable_text_characters_n": len(extracted),
                    **keyword_hits(extracted),
                }
            )
            log(
                f"REPOSITORY {pmcid} generic={generic} status={status} "
                f"url={link}"
            )

    supplements = pd.DataFrame(supplement_rows)
    repositories = pd.DataFrame(repo_rows)
    if supplements.empty:
        supplements = pd.DataFrame(
            columns=["pmcid", "supplement_href", "http_status", "extraction_status"]
        )
    if repositories.empty:
        repositories = pd.DataFrame(
            columns=[
                "pmcid",
                "repository_url",
                "generic_repository",
                "http_status",
                "extraction_status",
            ]
        )

    scope_rows: list[dict[str, object]] = []
    for _, article in sample.iterrows():
        pmcid = str(article["pmcid"])
        links = sorted(repo_by_pmc.get(pmcid, set()))
        article_links = [link for link in links if not is_generic_repo(link)]
        generic_links = [link for link in links if is_generic_repo(link)]
        repo_records = repositories.loc[
            repositories["pmcid"].astype(str).eq(pmcid)
            & ~repositories["generic_repository"].astype(bool)
        ]
        retrieved_repos = int(repo_records["http_status"].eq(200).sum())
        linked_supp = linked_supplements_by_pmc[pmcid]
        retrieved_supp = retrieved_supplements_by_pmc[pmcid]
        if linked_supp == 0:
            supplement_status = "none_linked"
        elif retrieved_supp == linked_supp:
            supplement_status = "retrieved_and_reviewed"
        elif retrieved_supp > 0:
            supplement_status = "partly_retrieved"
        else:
            supplement_status = "linked_unavailable"
        if article_links:
            repo_status = (
                "retrieved_and_reviewed"
                if retrieved_repos == len(article_links)
                else "linked_unavailable"
            )
        elif generic_links:
            repo_status = "generic_repository_only"
        else:
            repo_status = "none_linked"
        scope_rows.append(
            {
                "sample_order": article["sample_order"],
                "pmid": article["pmid"],
                "pmcid": pmcid,
                "doi": article["doi"],
                "title": article["title"],
                "main_text_reviewed": True,
                "supplements_linked_n": linked_supp,
                "supplements_retrieved_n": retrieved_supp,
                "supplements_text_reviewable_n": reviewable_supplements_by_pmc[pmcid],
                "supplement_status": supplement_status,
                "article_specific_repositories_linked_n": len(article_links),
                "article_specific_repositories_retrieved_n": retrieved_repos,
                "generic_repository_links_n": len(generic_links),
                "article_specific_repo_status": repo_status,
                "article_specific_repository_urls": "|".join(article_links),
                "generic_repository_urls": "|".join(generic_links),
            }
        )
    scope = pd.DataFrame(scope_rows)

    worksheet = sample[
        [
            "sample_order",
            "random_rank",
            "pmid",
            "pmcid",
            "doi",
            "publication_year",
            "title",
        ]
    ].copy()
    worksheet = worksheet.merge(
        scope[
            [
                "pmcid",
                "supplement_status",
                "article_specific_repo_status",
                "article_specific_repository_urls",
            ]
        ],
        on="pmcid",
        how="left",
    )
    for column in (
        "coder_id",
        "review_date",
        "main_text_reviewed",
        "supplement_reviewed",
        "linked_repo_reviewed",
        "source_layer",
        "named_native_table_reported",
        "database_identity_rule_reported",
        "time_origin_and_window_reported",
        "event_semantics_reported",
        "dose_or_route_reported",
        "evidence_location",
        "coder_notes",
    ):
        worksheet[column] = ""

    summary = pd.DataFrame(
        [
            {
                "metric": "sampled studies",
                "n": 40,
                "denominator": 40,
                "percent": 100.0,
            },
            {
                "metric": "main texts reviewed",
                "n": int(scope["main_text_reviewed"].sum()),
                "denominator": 40,
                "percent": 100.0 * float(scope["main_text_reviewed"].mean()),
            },
            {
                "metric": "studies with linked supplements",
                "n": int(scope["supplements_linked_n"].gt(0).sum()),
                "denominator": 40,
                "percent": 100.0 * float(scope["supplements_linked_n"].gt(0).mean()),
            },
            {
                "metric": "studies with all linked supplements retrieved",
                "n": int(scope["supplement_status"].eq("retrieved_and_reviewed").sum()),
                "denominator": 40,
                "percent": 100.0
                * float(scope["supplement_status"].eq("retrieved_and_reviewed").mean()),
            },
            {
                "metric": "studies with article-specific linked repositories",
                "n": int(scope["article_specific_repositories_linked_n"].gt(0).sum()),
                "denominator": 40,
                "percent": 100.0
                * float(scope["article_specific_repositories_linked_n"].gt(0).mean()),
            },
            {
                "metric": "article-specific linked repositories retrieved",
                "n": int(scope["article_specific_repositories_retrieved_n"].sum()),
                "denominator": int(scope["article_specific_repositories_linked_n"].sum()),
                "percent": (
                    100.0
                    * float(scope["article_specific_repositories_retrieved_n"].sum())
                    / float(scope["article_specific_repositories_linked_n"].sum())
                    if int(scope["article_specific_repositories_linked_n"].sum())
                    else 0.0
                ),
            },
        ]
    )

    flow_limit = pd.DataFrame(
        [
            {
                "stage": "PubMed records returned by frozen query",
                "n": 379,
                "interpretation": "frozen search universe",
            },
            {
                "stage": "English non-review records with open full text",
                "n": 293,
                "interpretation": "eligible sampling frame",
            },
            {
                "stage": "Records excluded by open-full-text and related filters",
                "n": 86,
                "interpretation": "direction of reporting-quality bias is unknown",
            },
        ]
    )

    outputs = {
        "supplement_inventory": TABLES / "published_operator_supplement_inventory.csv",
        "repository_inventory": TABLES / "published_operator_repository_inventory.csv",
        "evidence_scope": TABLES / "published_operator_evidence_scope.csv",
        "evidence_scope_summary": TABLES / "published_operator_evidence_scope_summary.csv",
        "blinded_recode_worksheet": TABLES / "published_operator_blinded_recode_worksheet.csv",
        "open_access_limitation": TABLES / "published_operator_open_access_flow_limitation.csv",
        "published_codebook_fields": TABLES / "published_operator_codebook_fields.csv",
    }
    supplements.to_csv(outputs["supplement_inventory"], index=False, encoding="utf-8-sig")
    repositories.to_csv(outputs["repository_inventory"], index=False, encoding="utf-8-sig")
    scope.to_csv(outputs["evidence_scope"], index=False, encoding="utf-8-sig")
    summary.to_csv(outputs["evidence_scope_summary"], index=False, encoding="utf-8-sig")
    worksheet.to_csv(outputs["blinded_recode_worksheet"], index=False, encoding="utf-8-sig")
    flow_limit.to_csv(outputs["open_access_limitation"], index=False, encoding="utf-8-sig")
    pd.read_csv(CODEBOOK_FIELDS).to_csv(
        outputs["published_codebook_fields"], index=False, encoding="utf-8-sig"
    )

    manifest = {
        "created_at": datetime.now().astimezone().isoformat(),
        "script": Path(__file__).name,
        "script_sha256": sha256(Path(__file__)),
        "frozen_hashes": {
            label: sha256(path) for label, (path, _) in hashes.items()
        },
        "sample_n": 40,
        "supplement_links_n": len(supplements),
        "supplements_http_200_n": int(supplements["http_status"].eq(200).sum()),
        "article_specific_repository_links_n": int(
            scope["article_specific_repositories_linked_n"].sum()
        ),
        "article_specific_repositories_http_200_n": int(
            scope["article_specific_repositories_retrieved_n"].sum()
        ),
        "outputs": {
            name: {
                "path": str(path.relative_to(PROJECT)),
                "sha256": sha256(path),
            }
            for name, path in outputs.items()
        },
    }
    (MANIFESTS / "37_published_operator_evidence_scope_manifest.json").write_text(
        json.dumps(manifest, indent=2), encoding="utf-8"
    )
    (ENVIRONMENT / "sessionInfo_published_operator_scope.txt").write_text(
        "\n".join(
            [
                f"timestamp={manifest['created_at']}",
                f"python={sys.version}",
                f"pandas={pd.__version__}",
                f"script={Path(__file__).name}",
                f"script_sha256={manifest['script_sha256']}",
                f"codebook_sha256={EXPECTED_CODEBOOK_SHA256}",
                f"sample_sha256={EXPECTED_SAMPLE_SHA256}",
            ]
        )
        + "\n",
        encoding="utf-8",
    )
    log("PASS published-operator evidence-scope audit complete")
    print(summary.to_string(index=False), flush=True)
    print(json.dumps(manifest, indent=2), flush=True)


if __name__ == "__main__":
    main()
