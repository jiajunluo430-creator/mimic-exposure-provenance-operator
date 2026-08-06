#!/usr/bin/env python3
"""Build and validate editable JBI submission documents from committed Markdown.

The script is deliberately dependency-light (python-docx only), fail-closed on
JBI count/citation rules, and writes no patient-level data.
"""

from __future__ import annotations

import csv
import json
import re
import shutil
import sys
from collections import OrderedDict
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_ROW_HEIGHT_RULE
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SRC = ROOT / "manuscript" / "jbi" / "source"
OUT = ROOT / "manuscript" / "jbi" / "submission_package"
FIG = ROOT / "manuscript" / "jbi" / "figure_redesign_2026-08-06_v3" / "final_figures"
GRAPHICAL_ABSTRACT = ROOT / "manuscript" / "jbi" / "graphical_abstract_2026-08-06"

NAVY = "17324D"
TEAL = "1E6F74"
MID = "52667A"
LIGHT = "EAF1F4"
WHITE = "FFFFFF"
BLACK = "1F2933"


def read_text(name: str) -> str:
    return (SRC / name).read_text(encoding="utf-8")


def strip_md(text: str) -> str:
    text = re.sub(r"`([^`]+)`", r"\1", text)
    text = re.sub(r"\*\*([^*]+)\*\*", r"\1", text)
    text = re.sub(r"\[([^]]+)\]\(([^)]+)\)", r"\1", text)
    text = text.replace("<!-- PAGEBREAK -->", "")
    return text


def word_count(text: str) -> int:
    clean = strip_md(text)
    clean = re.sub(r"^#+\s*", "", clean, flags=re.MULTILINE)
    clean = re.sub(r"^\s*[-*]\s+", "", clean, flags=re.MULTILINE)
    clean = re.sub(r"^\s*\d+[.)]\s+", "", clean, flags=re.MULTILINE)
    clean = re.sub(r"\|[-: ]+", " ", clean)
    return len(re.findall(r"\b[\w’'-]+\b", clean, flags=re.UNICODE))


def between(text: str, start: str, end: str) -> str:
    i = text.index(start) + len(start)
    j = text.index(end, i)
    return text[i:j]


def expand_citation_group(group: str) -> list[int]:
    values: list[int] = []
    for item in re.split(r"\s*,\s*", group):
        item = item.strip()
        if not item:
            continue
        m = re.fullmatch(r"(\d+)\s*[–-]\s*(\d+)", item)
        if m:
            a, b = map(int, m.groups())
            values.extend(range(a, b + 1))
        elif item.isdigit():
            values.append(int(item))
    return values


def validate_sources() -> dict:
    main = read_text("JBI_main_manuscript.md")
    abstract = between(main, "# Abstract", "**Keywords:**")
    body = between(main, "# 1. Introduction", "# CRediT authorship contribution statement")
    significance = between(
        main, "## 1.1. Statement of significance", "# 2. Related work"
    )
    refs = main.split("# References", 1)[1]
    ref_ids = [int(x) for x in re.findall(r"(?m)^(\d+)\.\s", refs)]
    if ref_ids != list(range(1, 43)):
        raise SystemExit(f"Reference list must be contiguous 1-42; got {ref_ids[:5]}...{ref_ids[-5:]}")

    pre_refs = main.split("# References", 1)[0]
    if re.search(r"\[\d", abstract):
        raise SystemExit("Numbered citations are not allowed in the abstract")
    cited: list[int] = []
    first_seen: list[int] = []
    for group in re.findall(r"\[([0-9,\-–\s]+)\]", pre_refs):
        for value in expand_citation_group(group):
            cited.append(value)
            if value not in first_seen:
                first_seen.append(value)
    if first_seen != list(range(1, 43)):
        raise SystemExit(f"First citation appearance must be 1-42; got {first_seen}")
    if set(cited) != set(ref_ids):
        raise SystemExit(f"Citation/list mismatch: cited={sorted(set(cited))}, refs={ref_ids}")

    abstract_wc = word_count(abstract)
    body_wc = word_count(body)
    if abstract_wc > 250:
        raise SystemExit(f"Abstract exceeds JBI limit: {abstract_wc}")
    if body_wc > 6000:
        raise SystemExit(f"Body exceeds JBI limit: {body_wc}")

    significance_rows = []
    for line in significance.splitlines():
        if not line.strip().startswith("|") or is_separator(line):
            continue
        cells = [cell.strip() for cell in line.strip().strip("|").split("|")]
        if len(cells) == 2 and cells[0] != "Statement element":
            significance_rows.append(cells)
    required_significance = [
        "Problem or Issue",
        "What is Already Known",
        "What this Paper Adds",
        "Who would benefit from the new knowledge in this paper",
    ]
    if [row[0] for row in significance_rows] != required_significance:
        raise SystemExit("Statement of significance must use the four JBI headings")
    significance_wc = sum(word_count(row[1]) for row in significance_rows)
    if significance_wc > 150:
        raise SystemExit(
            f"Statement of significance exceeds JBI limit: {significance_wc}"
        )

    main_tables = len(re.findall(r"(?m)^## Table \d+\.", main))
    main_figures = len(re.findall(r"(?m)^## Figure \d+\.", main))
    if main_tables + main_figures > 8:
        raise SystemExit(f"JBI display limit exceeded: {main_tables}+{main_figures}")

    highlights = [strip_md(x[2:].strip()) for x in read_text("JBI_highlights.md").splitlines() if x.startswith("- ")]
    if not (3 <= len(highlights) <= 5):
        raise SystemExit(f"Highlights must contain 3-5 bullets; got {len(highlights)}")
    over = [(h, len(h)) for h in highlights if len(h) > 85]
    if over:
        raise SystemExit(f"Highlights exceed 85 characters: {over}")

    with (SRC / "reference_registry.csv").open(newline="", encoding="utf-8") as handle:
        registry = list(csv.DictReader(handle))
    if len(registry) != 42 or [int(row["reference_id"]) for row in registry] != list(range(1, 43)):
        raise SystemExit("Reference registry must contain ordered IDs 1-42")
    if any(row["status"] != "VERIFIED" for row in registry):
        raise SystemExit("Reference registry contains an unverified entry")

    return {
        "abstract_words": abstract_wc,
        "main_text_words": body_wc,
        "statement_of_significance_words": significance_wc,
        "references": len(ref_ids),
        "main_tables": main_tables,
        "main_figures": main_figures,
        "main_displays": main_tables + main_figures,
        "highlights": [{"text": h, "characters": len(h)} for h in highlights],
        "citations_first_appearance": first_seen,
        "reference_registry_verified": len(registry),
    }


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=70, start=80, bottom=70, end=80) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table) -> None:
    """Apply a restrained three-line journal table with no vertical rules."""

    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for child in list(borders):
        borders.remove(child)
    for edge, value, size, color in (
        ("top", "single", "10", NAVY),
        ("bottom", "single", "10", NAVY),
        ("left", "nil", "0", WHITE),
        ("right", "nil", "0", WHITE),
        ("insideH", "nil", "0", WHITE),
        ("insideV", "nil", "0", WHITE),
    ):
        node = OxmlElement(f"w:{edge}")
        node.set(qn("w:val"), value)
        node.set(qn("w:sz"), size)
        node.set(qn("w:color"), color)
        borders.append(node)


def set_cell_bottom_border(cell, color: str = MID) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.find(qn("w:tcBorders"))
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    bottom = borders.find(qn("w:bottom"))
    if bottom is None:
        bottom = OxmlElement("w:bottom")
        borders.append(bottom)
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "8")
    bottom.set(qn("w:color"), color)


def keep_with_next(paragraph, value: bool = True) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    node = p_pr.find(qn("w:keepNext"))
    if value and node is None:
        node = OxmlElement("w:keepNext")
        p_pr.append(node)


def prevent_row_split(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr_text)
    run._r.append(fld_char2)


def configure_document(doc: Document, title: str, kind: str) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.85)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.35)

    if kind == "cover":
        section.top_margin = Inches(0.65)
        section.bottom_margin = Inches(0.65)
        section.left_margin = Inches(0.85)
        section.right_margin = Inches(0.85)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5 if kind in {"support", "cover"} else 11)
    normal.font.color.rgb = RGBColor.from_string(BLACK)
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    normal.paragraph_format.line_spacing = 1.15

    for idx, size, color in ((1, 14, NAVY), (2, 12, TEAL), (3, 11, MID)):
        style = styles[f"Heading {idx}"]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(10 if idx == 1 else 7)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.keep_with_next = True

    title_style = styles["Title"]
    title_style.font.name = "Calibri"
    title_style.font.size = Pt(17)
    title_style.font.bold = True
    title_style.font.color.rgb = RGBColor.from_string(NAVY)
    title_style.paragraph_format.space_after = Pt(12)

    for section in doc.sections:
        if kind not in {"cover", "highlights"}:
            header = section.header.paragraphs[0]
            header.text = "EXECUTABLE MEDICATION-EXPOSURE PROVENANCE"
            header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            for run in header.runs:
                run.font.name = "Calibri"
                run.font.size = Pt(8)
                run.font.color.rgb = RGBColor.from_string(MID)
            add_page_number(section.footer.paragraphs[0])

    doc.core_properties.title = title
    doc.core_properties.subject = "Journal of Biomedical Informatics submission"
    doc.core_properties.author = "Jiajun Luo; Qinglong Chen; Jing Liu; Fanghui Lu; Xiaolong Liang"
    doc.core_properties.keywords = "medication exposure; provenance; EHR; FHIR; OMOP"


INLINE_TOKEN = re.compile(r"(\*\*[^*]+\*\*|`[^`]+`|https?://\S+)")


def add_inline(paragraph, text: str) -> None:
    text = text.replace("  \n", "\n")
    parts = INLINE_TOKEN.split(text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            run.font.name = "Consolas"
            run.font.size = Pt(9.2)
            run.font.color.rgb = RGBColor.from_string(TEAL)
        elif part.startswith("http://") or part.startswith("https://"):
            run = paragraph.add_run(part)
            run.font.color.rgb = RGBColor(0x05, 0x63, 0xC1)
            run.underline = True
        else:
            subparts = part.split("\n")
            for i, sub in enumerate(subparts):
                if i:
                    paragraph.add_run().add_break()
                paragraph.add_run(sub)


def add_table(doc: Document, rows: list[list[str]], kind: str) -> None:
    if not rows:
        return
    ncols = max(len(row) for row in rows)
    table = doc.add_table(rows=len(rows), cols=ncols)
    table.style = "Table Grid"
    set_table_borders(table)
    table.autofit = True
    table.alignment = 1
    font_size = 7.4 if ncols >= 7 else (8.0 if kind == "support" else 8.3)
    for r_idx, source_row in enumerate(rows):
        row = table.rows[r_idx]
        prevent_row_split(row)
        if r_idx == 0:
            repeat_table_header(row)
        for c_idx in range(ncols):
            cell = row.cells[c_idx]
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
            value = source_row[c_idx].strip() if c_idx < len(source_row) else ""
            cell.text = ""
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.0
            add_inline(p, value)
            # Keep the header row with at least the first data row. Without
            # this, Word can leave a repeated table header orphaned at the
            # bottom of a page while moving every data row to the next page.
            if r_idx == 0 or (kind == "main" and r_idx < len(rows) - 1):
                keep_with_next(p)
            for run in p.runs:
                run.font.name = "Calibri"
                run.font.size = Pt(font_size)
                if r_idx == 0:
                    run.bold = True
                    run.font.color.rgb = RGBColor.from_string(NAVY)
            if r_idx == 0:
                set_cell_bottom_border(cell)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def is_separator(line: str) -> bool:
    cells = [x.strip() for x in line.strip().strip("|").split("|")]
    return bool(cells) and all(re.fullmatch(r":?-{3,}:?", cell) for cell in cells)


def parse_table(lines: list[str], start: int) -> tuple[list[list[str]], int]:
    rows: list[list[str]] = []
    i = start
    while i < len(lines) and lines[i].strip().startswith("|"):
        if not is_separator(lines[i]):
            rows.append([x.strip() for x in lines[i].strip().strip("|").split("|")])
        i += 1
    return rows, i


def build_doc(markdown: str, out_path: Path, kind: str, stats: dict) -> None:
    markdown = markdown.replace("to be populated from validated build", "{COUNT}")
    if kind == "title":
        markdown = markdown.replace("- Abstract: {COUNT}", f"- Abstract: {stats['abstract_words']} words")
        markdown = markdown.replace("- Main text (Introduction through Conclusion): {COUNT}", f"- Main text (Introduction through Conclusion): {stats['main_text_words']} words")

    lines = markdown.splitlines()
    first_title = next((strip_md(line[2:].strip()) for line in lines if line.startswith("# ")), out_path.stem)
    doc = Document()
    configure_document(doc, first_title, kind)

    i = 0
    title_done = False
    before_first_section = True
    while i < len(lines):
        raw = lines[i]
        line = raw.strip()
        if not line:
            i += 1
            continue
        if line == "<!-- PAGEBREAK -->":
            doc.add_page_break()
            before_first_section = False
            i += 1
            continue
        if line.startswith("|") and i + 1 < len(lines) and is_separator(lines[i + 1]):
            rows, i = parse_table(lines, i)
            add_table(doc, rows, kind)
            continue
        heading = re.match(r"^(#{1,3})\s+(.+)$", line)
        if heading:
            level = len(heading.group(1))
            text = strip_md(heading.group(2))
            if not title_done:
                p = doc.add_paragraph(style="Title")
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                add_inline(p, heading.group(2))
                title_done = True
                before_first_section = True
            else:
                if kind == "main" and level == 1 and text in {"Tables", "Figure legends", "References"}:
                    doc.add_page_break()
                p = doc.add_paragraph(style=f"Heading {min(level, 3)}")
                add_inline(p, heading.group(2))
                keep_with_next(p)
                if level == 1:
                    before_first_section = False
            i += 1
            continue
        if re.match(r"^[-*]\s+", line):
            p = doc.add_paragraph(style="List Bullet")
            add_inline(p, re.sub(r"^[-*]\s+", "", line))
            p.paragraph_format.space_after = Pt(2)
            i += 1
            continue
        if re.match(r"^\d+[.)]\s+", line):
            p = doc.add_paragraph(style="List Number")
            add_inline(p, re.sub(r"^\d+[.)]\s+", "", line))
            p.paragraph_format.space_after = Pt(2)
            i += 1
            continue

        para_lines = [raw.rstrip()]
        i += 1
        while i < len(lines):
            nxt = lines[i]
            stripped = nxt.strip()
            if not stripped or stripped == "<!-- PAGEBREAK -->" or re.match(r"^#{1,3}\s+", stripped) or re.match(r"^[-*]\s+", stripped) or re.match(r"^\d+[.)]\s+", stripped):
                break
            if stripped.startswith("|") and i + 1 < len(lines) and is_separator(lines[i + 1]):
                break
            para_lines.append(nxt.rstrip())
            i += 1
        text = "\n".join(x[:-2] if x.endswith("  ") else x for x in para_lines)
        p = doc.add_paragraph()
        add_inline(p, text)
        if kind in {"main", "title"} and before_first_section:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if kind == "cover":
            p.paragraph_format.line_spacing = 1.02
            p.paragraph_format.space_after = Pt(4)
            p.paragraph_format.keep_together = True
        if kind == "highlights":
            p.paragraph_format.space_after = Pt(5)

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out_path)


def copy_figures() -> list[str]:
    copied: list[str] = []
    for idx in range(1, 6):
        for suffix in (".pdf", ".tiff"):
            src = FIG / f"JBI_Figure{idx}{suffix}"
            if not src.exists():
                raise SystemExit(f"Missing required figure: {src}")
            dst = OUT / f"JBI_Figure{idx}{suffix}"
            shutil.copy2(src, dst)
            copied.append(dst.name)
    for suffix in (".pdf", ".tiff", ".svg"):
        src = GRAPHICAL_ABSTRACT / f"JBI_Graphical_Abstract{suffix}"
        if not src.exists():
            raise SystemExit(f"Missing required graphical abstract: {src}")
        dst = OUT / src.name
        shutil.copy2(src, dst)
        copied.append(dst.name)
    return copied


def main() -> int:
    OUT.mkdir(parents=True, exist_ok=True)
    stats = validate_sources()
    jobs = OrderedDict(
        [
            ("JBI_main_manuscript.md", ("JBI_main_manuscript.docx", "main")),
            ("JBI_title_page.md", ("JBI_title_page.docx", "title")),
            ("JBI_cover_letter.md", ("JBI_cover_letter.docx", "cover")),
            ("JBI_highlights.md", ("JBI_highlights.docx", "highlights")),
            (
                "JBI_declaration_of_competing_interest.md",
                ("JBI_declaration_of_competing_interest.docx", "highlights"),
            ),
            ("JBI_supporting_information.md", ("JBI_supporting_information.docx", "support")),
            ("JBI_RECORD_STROBE_checklist.md", ("JBI_RECORD_STROBE_checklist.docx", "support")),
        ]
    )
    for source_name, (output_name, kind) in jobs.items():
        build_doc(read_text(source_name), OUT / output_name, kind, stats)
    stats["figure_files"] = copy_figures()
    stats["documents"] = [x[0] for x in jobs.values()]
    stats["status"] = "PASS_JBI_SOURCE_AND_DOCUMENT_BUILD"
    (OUT / "JBI_BUILD_VALIDATION.json").write_text(json.dumps(stats, indent=2, ensure_ascii=False) + "\n", encoding="utf-8")
    report = [
        "# JBI build validation",
        "",
        "**PASS_JBI_SOURCE_AND_DOCUMENT_BUILD**",
        "",
        f"- Abstract: {stats['abstract_words']} words (limit 250).",
        f"- Main text: {stats['main_text_words']} words (limit 6000).",
        f"- Statement of significance: {stats['statement_of_significance_words']} words (limit 150; four required headings present).",
        f"- References: {stats['references']}; citation order 1–42 passed.",
        f"- Main displays: {stats['main_figures']} figures + {stats['main_tables']} tables = {stats['main_displays']} (limit 8).",
        f"- Highlights: {len(stats['highlights'])}; all ≤85 characters.",
        "- Reference registry: 42/42 entries verified.",
        "- Editable DOCX files and separate PDF/TIFF figures generated.",
        "",
    ]
    (OUT / "JBI_BUILD_VALIDATION.md").write_text("\n".join(report), encoding="utf-8")
    print(json.dumps(stats, indent=2, ensure_ascii=False))
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
